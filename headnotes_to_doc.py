import pandas as pd
import re
from datetime import date
from docx import Document

# Lade die CSV-Datei
df = pd.read_csv("ep_appeal_decisions.csv", encoding="utf-8")

# Direkt filtern nach FP1, EP1, DP1 usw. (jede Zahl hinter FP, EP, DP)
df = df[df["Reference"].str.contains(r'(FP\d+|EP\d+|DP\d+)', regex=True, na=False)].copy()

# Funktion zur Extraktion von Sprache und Typ
def extract_language_and_type(reference):
    match = re.match(r'([A-Z]{1,2})(\d+)([A-Z]+)', reference)
    if match:
        #print(match.group(3))
        lang_code = match.group(3)
        decision_type = match.group(1)[0]  # Nur den ersten Buchstaben des Codes (z. B. 'F' -> 'F' für Französisch)
        
        # Mapping für die Sprache
        language_map = {
            "FP": "French", "EP": "English", "DP": "German",
        }
        language = language_map.get(lang_code, "Unknown")
        
        return language, decision_type
    return "Unknown", "Unknown"

# Extrahiere Sprache und Typ
df["Language"], df["DecisionType"] = zip(*df["Reference"].apply(extract_language_and_type))

# Heutiges Datum
date_today = date.today().strftime("%Y-%m-%d")

# Anzahl extrahierter Entscheidungen
extrahiert = len(df)

# Reihenfolge für die Sprachen und Entscheidungstypen
language_order = ["English", "French", "German"]
decision_type_order = ["G", "J", "T", "D", "W"]

# Sortiere nach der festgelegten Reihenfolge der Sprache und des Entscheidungstyps
df_sorted = df[df["Language"].isin(language_order)].copy()
df_sorted["Language"] = pd.Categorical(df_sorted["Language"], categories=language_order, ordered=True)
df_sorted["DecisionType"] = pd.Categorical(df_sorted["DecisionType"], categories=decision_type_order, ordered=True)
df_sorted = df_sorted.sort_values(by=["Language", "DecisionType"])

# Gemeinsamer Inhalt für alle Dateien
header = (
    "# Headnotes of Decisions of the Boards of Appeal of the European Patent Office\n\n"
    f"**Date**: {date_today}\n"
    f"**Number of Headnotes**: {extrahiert}\n\n"
    "## Source\n"
    "14.6 EPO Boards of Appeal decisions - BoA decisions September 2024 - "
    "https://publication-bdds.apps.epo.org/raw-data/products/public/product/21\n\n"
    "## Decisions\n\n"
)

# Markdown-Datei erstellen
def create_md_file(filename="epo_decisions.md"):
    with open(filename, "w", encoding="utf-8") as file:
        file.write(header)
        
        # Gruppiere nach Sprache und Entscheidungstyp und schreibe die entsprechenden Überschriften
        for language in language_order:
            file.write(f"## {language}\n")
            for decision_type in decision_type_order:
                group = df_sorted[(df_sorted["Language"] == language) & (df_sorted["DecisionType"] == decision_type)]
                if not group.empty:
                    file.write(f"### {decision_type}-Decisions\n")
                    for _, row in group.iterrows():
                        if pd.notna(row["Reference"]) and pd.notna(row["ep-headnote"]):
                            file.write(f"#### {row['Reference']}\n")
                            file.write(f"{row['ep-headnote']}\n\n")

# TXT-Datei erstellen
def create_txt_file(filename="epo_decisions.txt"):
    with open(filename, "w", encoding="utf-8") as file:
        file.write(header.replace("#", "").replace("##", ""))  # Entferne Markdown-Überschriften
        
        # Gruppiere nach Sprache und Entscheidungstyp und schreibe die entsprechenden Überschriften
        for language in language_order:
            file.write(f"{language}\n")
            for decision_type in decision_type_order:
                group = df_sorted[(df_sorted["Language"] == language) & (df_sorted["DecisionType"] == decision_type)]
                if not group.empty:
                    file.write(f"{decision_type}-Decisions\n")
                    for _, row in group.iterrows():
                        if pd.notna(row["Reference"]) and pd.notna(row["ep-headnote"]):
                            file.write(f"{row['Reference']}\n")
                            file.write(f"{row['ep-headnote']}\n\n")

# DOCX-Datei erstellen
def create_docx_file(filename="epo_decisions.docx"):
    doc = Document()
    doc.add_heading("Headnotes of Decisions of the Boards of Appeal of the European Patent Office", level=1)
    doc.add_paragraph(f"Date: {date_today}")
    doc.add_paragraph(f"Number of Headnotes: {extrahiert}\n")
    doc.add_heading("Source", level=2)
    doc.add_paragraph("14.6 EPO Boards of Appeal decisions - BoA decisions September 2024 - "
                      "https://publication-bdds.apps.epo.org/raw-data/products/public/product/21")
    doc.add_heading("Decisions", level=2)

    # Gruppiere nach Sprache und Entscheidungstyp und schreibe die entsprechenden Überschriften
    for language in language_order:
        doc.add_heading(language, level=2)
        for decision_type in decision_type_order:
            group = df_sorted[(df_sorted["Language"] == language) & (df_sorted["DecisionType"] == decision_type)]
            if not group.empty:
                doc.add_heading(f"{decision_type}-Decisions", level=3)
                for _, row in group.iterrows():
                    if pd.notna(row["Reference"]) and pd.notna(row["ep-headnote"]):
                        doc.add_heading(row["Reference"], level=4)
                        doc.add_paragraph(row["ep-headnote"])

    doc.save(filename)

# Dateien erstellen
create_md_file()
create_txt_file()
create_docx_file()

print(f"Markdown-, TXT- und DOCX-Dateien mit {extrahiert} Entscheidungen wurden erstellt.")
# Statistik

# Funktion zum Extrahieren der Sprache (Zahl am Ende ignorieren)
def extract_language(reference):
    match = re.search(r'([A-Z]{1,2})\d+$', reference)  # Nur den Buchstaben-Teil extrahieren
    if match:
        lang_code = match.group(1)
        language_map = {
            "FP": "FR", "EP": "EN", "DP": "DE",  
            # "FX": "FR (Verfahrenssprache)", "EX": "EN (Verfahrenssprache)", "DX": "DE (Verfahrenssprache)",
            # "FU": "FR (nicht veröffentlicht)", "EU": "EN (nicht veröffentlicht)", "DU": "DE (nicht veröffentlicht)",
        }
        return language_map.get(lang_code, "Unknown")
    return "Unknown"

# Funktion zum Extrahieren des Entscheidungstyps
def extract_decision_type(reference):
    return reference[0]  # Der erste Buchstabe gibt den Typ an

# Sprache und Typ extrahieren
df["Language"] = df["Reference"].apply(extract_language)
df["Decision Type"] = df["Reference"].apply(extract_decision_type)

# Gruppieren nach Sprache und Typ
language_type_counts = df.groupby(["Language", "Decision Type"]).size().unstack(fill_value=0)

# Gruppieren nach Sprache und Typ
total_decisions_per_type = df["Decision Type"].value_counts()

# Berechnung der Gesamtzahl der Entscheidungen pro Sprache
total_decisions_per_language = df["Language"].value_counts()

# Ausgabe der Ergebnisse
print("Entscheidungen pro Sprache und Typ:")
print(language_type_counts)
print("\nGesamtzahl der Entscheidungen pro Typ:")
print(total_decisions_per_type)
print("\nGesamtzahl der Entscheidungen pro Sprache:")
print(total_decisions_per_language)

# Speichern als Excel-Datei
with pd.ExcelWriter("decision_analysis.xlsx") as writer:
    language_type_counts.to_excel(writer, sheet_name="Nach Sprache & Typ")
    total_decisions_per_language.to_excel(writer, sheet_name="Gesamt pro Sprache")

           


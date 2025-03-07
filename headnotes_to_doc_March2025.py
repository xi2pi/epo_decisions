import pandas as pd
import re
from datetime import date
from docx import Document

# Lade die CSV-Datei mit korrekter Kodierung
df = pd.read_csv("ep_appeal_decisions.csv", encoding='ISO-8859-1')
# g1=df[df["Reference"].str.startswith("G22")]
# print(g1)


# # Replace 'None' string with actual None (NaN) for proper handling
df['ep-catchword-language'] = df['ep-catchword-language'].replace('None', None)
df['ep-headnote-language'] = df['ep-headnote-language'].replace('None', None)

# # Merge the columns, taking the non-None value
df["Language"] = df['ep-catchword-language'].combine_first(df['ep-headnote-language'])
df["Decision Type"] =  df["Case Code"]

# Funktion zur Erstellung des Aktenzeichens
def generate_case_number(reference):
    match = re.match(r'([GJTDW])(\d{2})(\d{4})[A-Z]+\d+', reference)
    if match:
        decision_type = match.group(1)
        case_number = match.group(3).lstrip("0")  # Entferne führende Nullen
        year = match.group(2)
        return f"{decision_type} {case_number}/{year}"
    return reference  # Falls kein Treffer, behalte die originale Referenz

df["Case Number"] = df["Reference"].apply(generate_case_number)

# Heutiges Datum
date_today = date.today().strftime("%Y-%m-%d")

# Entferne Duplikate basierend auf Headnote oder Catchword
df = df.drop_duplicates(subset=["ep-headnote", "ep-catchword"], keep="first")

# Anzahl extrahierter Entscheidungen
extrahiert = len(df)

# Reihenfolge für die Sortierung
language_order = ["EN", "FR", "DE"]
decision_type_order = ["G", "J", "T", "D", "W"]

# Konvertiere das Jahr in numerischen Wert für die Sortierung
df["Year"] = pd.to_numeric(df["Year"], errors="coerce")

# Sortiere nach Sprache, Entscheidungstyp und Jahr (neueste zuerst)
df_sorted = df.sort_values(by=["Language", "Decision Type", "Year"], ascending=[True, True, False], 
                           key=lambda x: x.map(lambda y: language_order.index(y) if y in language_order else len(language_order)) if x.name == "Language" else 
                                       x.map(lambda y: decision_type_order.index(y) if y in decision_type_order else len(decision_type_order)) if x.name == "Decision Type" else x)

# g2=df[df["Reference"].str.startswith("G22")]
# print(g2)


# TXT-Datei erstellen
txt_filename = "epo_decisions.txt"
with open(txt_filename, "w", encoding="utf-8") as file:
    file.write(f"Headnotes and Catchwords of EPO Board of Appeal Decisions\n\n")
    file.write(f"Date: {date_today}\n")
    file.write(f"Total Extracted Decisions: {extrahiert}\n\n")

    for language in language_order:
        group_lang = df_sorted[df_sorted["Language"] == language]
        if not group_lang.empty:
            file.write(f"\n{language} Decisions\n")
            for decision_type in decision_type_order:
                group = group_lang[group_lang["Decision Type"] == decision_type]
                if not group.empty:
                    file.write(f"\n{decision_type}-Decisions\n")
                    for _, row in group.iterrows():
                        if pd.notna(row["Reference"]):
                            file.write(f"\n{row['Case Number']} ({row['Reference']})\n")
                            if pd.notna(row["ep-headnote"]):
                                file.write(f"Headnote: {row['ep-headnote']}\n")
                            if pd.notna(row["ep-catchword"]):
                                file.write(f"Catchword: {row['ep-catchword']}\n")

# MD-Datei erstellen
md_filename = "epo_decisions.md"
with open(md_filename, "w", encoding="utf-8") as file:
    file.write(f"# Headnotes and Catchwords of EPO Board of Appeal Decisions\n\n")
    file.write(f"**Date:** {date_today}\n")
    file.write(f"**Total Extracted Decisions:** {extrahiert}\n\n")

    for language in language_order:
        group_lang = df_sorted[df_sorted["Language"] == language]
        if not group_lang.empty:
            file.write(f"\n## {language} Decisions\n")
            for decision_type in decision_type_order:
                group = group_lang[group_lang["Decision Type"] == decision_type]
                if not group.empty:
                    file.write(f"\n### {decision_type}-Decisions\n")
                    for _, row in group.iterrows():
                        if pd.notna(row["Reference"]):
                            file.write(f"\n#### {row['Case Number']} ({row['Reference']})\n")
                            if pd.notna(row["ep-headnote"]):
                                file.write(f"**Headnote:** {row['ep-headnote']}\n")
                            if pd.notna(row["ep-catchword"]):
                                file.write(f"**Catchword:** {row['ep-catchword']}\n")

# DOCX-Datei erstellen
doc_filename = "epo_decisions.docx"
doc = Document()
doc.add_heading("Headnotes and Catchwords of EPO Board of Appeal Decisions", level=1)
doc.add_paragraph(f"Date: {date_today}")
doc.add_paragraph(f"Total Extracted Decisions: {extrahiert}\n")

for language in language_order:
    group_lang = df_sorted[df_sorted["Language"] == language]
    if not group_lang.empty:
        doc.add_heading(f"{language} Decisions", level=2)
        for decision_type in decision_type_order:
            group = group_lang[group_lang["Decision Type"] == decision_type]
            if not group.empty:
                doc.add_heading(f"{decision_type}-Decisions", level=3)
                for _, row in group.iterrows():
                    if pd.notna(row["Reference"]):
                        doc.add_heading(f"{row['Case Number']} ({row['Reference']})", level=4)
                        if pd.notna(row["ep-headnote"]):
                            doc.add_paragraph(f"Headnote: {row['ep-headnote']}")
                        if pd.notna(row["ep-catchword"]):
                            doc.add_paragraph(f"Catchword: {row['ep-catchword']}")

doc.save(doc_filename)

print(f"TXT, MD und DOCX wurden erfolgreich mit neuer Sortierung erstellt.")


# Statistik erstellen
language_type_counts = df.groupby(["Language", "Case Code"]).size().unstack(fill_value=0)
total_decisions_per_language = df["Language"].value_counts()

# Speichern als Excel-Datei
excel_filename = "decision_analysis.xlsx"
with pd.ExcelWriter(excel_filename) as writer:
    language_type_counts.to_excel(writer, sheet_name="Nach Sprache & Typ")
    total_decisions_per_language.to_excel(writer, sheet_name="Gesamt pro Sprache")

print(f"Statistik wurde als {excel_filename} gespeichert.")

